from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

doc.add_heading("ACD Online — What changed (Aug 3–4)", 0)
doc.add_paragraph("Quick notes. Edit as you like.")

doc.add_heading("SF2 / students", level=1)
for b in [
    "Uploaded students had grades/sections that weren’t always in School setup — they only lived on the student row.",
    "We now copy grade + section from students into School setup (also after import).",
    "SF2 section dropdown was stuck empty online because the page script often didn’t run on Hostinger.",
    "Section list now comes from students + setup. If nothing shows, you can type the section.",
    "“Load from attendance logs” was doing nothing for the same reason — should now say Loading… or show an alert.",
    "Note: load only fills learners who have grade, section, and sex set.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Attendance time", level=1)
doc.add_paragraph(
    "Tardy / expected login rules (and the policy settings page) got a few tweaks so attendance matches the school’s time rules better."
)

doc.add_heading("Hostinger / ops", level=1)
for b in [
    "Small helper commands for scheduler check and rolling back auto attendance scans.",
    "Hostinger deploy notes updated a bit.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("After pull on Hostinger", level=1)
for b in [
    "Pull main (need 317ba08 or later).",
    "php artisan view:clear",
    "Hard-refresh SF2 Create (Ctrl+F5).",
    "Try: grade → section → Load from attendance logs.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Still to do", level=1)
doc.add_paragraph("Things not finished yet / worth handling next:")
for b in [
    "Deploy & verify on Hostinger — push/pull latest main, clear views, hard-refresh, confirm SF2 section list + Load from logs both work online.",
    "Clean student data — give Kinder (and anyone else) a section; set sex (male/female) so SF2 can load rosters.",
    "Check production grades — open School setup after deploy and confirm real sections listed (not only “Test”); fix odd grade labels if any rows still use “9” / “G9” etc.",
    "Fill missing years on SF2 — grades with no students/sections (e.g. G9 if empty) need either import or manual sections in School setup.",
    "Training note for staff — short reminder: grade → section → Load from logs; hard-refresh once after site updates.",
    "Optional: smoke-test attendance times/tardy after the policy tweaks (one sample class, one day).",
    "Optional: confirm Hostinger cron is actually firing (scheduler ping / daily attendance jobs).",
    "Optional: document when to use RollbackAutoAttendanceScans so ops isn’t guessing later.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Commits (newest first)", level=1)
for b in [
    "317ba08 — SF2 dropdown + load-from-logs",
    "944d0cd / 2ef87f4 — more SF2 / section sync",
    "97f6557 — scheduler / rollback tools",
    "4ce6d57 / ec91345 — attendance time / sections policy",
]:
    doc.add_paragraph(b, style="List Bullet")

out = Path(r"d:\acd_online\docs\Change_Summary_Aug3-4_short.docx")
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(out)
